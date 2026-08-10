#!/usr/bin/env python3
"""
Build the final LuxeStay thesis DOCX from markdown content + screenshots.

Prerequisites:
  pip install python-docx Pillow

Usage:
  python docs/tools/build_thesis_docx.py

Output:
  docs/thesis-final/LuxeStay_KhoaLuan_FINAL_v2.docx
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# Paths
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
SCREENSHOT_DIR = PROJECT_ROOT / "docs" / "screenshots" / "thesis-final"
DIAGRAM_DIR = PROJECT_ROOT / "docs" / "thesis-assets" / "diagrams"
OUTPUT_DIR = PROJECT_ROOT / "docs" / "thesis-final"
OUTPUT_FILE = OUTPUT_DIR / "LuxeStay_KhoaLuan_FINAL_v2.docx"


def setup_document():
    """Create a new document with proper formatting according to Mau 5."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Paragraph format
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # Set margins: top=2cm, bottom=2cm, left=3cm, right=2cm
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Setup Heading styles
    for level in range(1, 4):
        style_name = f'Heading {level}'
        heading_style = doc.styles[style_name]
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(13)

    return doc


def add_cover_page(doc):
    """Add the thesis cover page."""
    # Header lines
    for text in [
        "TRƯỜNG ĐẠI HỌC SƯ PHẠM KỸ THUẬT TP.HCM",
        "KHOA CÔNG NGHỆ THÔNG TIN",
        "BỘ MÔN CÔNG NGHỆ PHẦN MỀM"
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.bold = True

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TIỂU LUẬN TỐT NGHIỆP KỸ SƯ CÔNG NGHỆ THÔNG TIN")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỀ TÀI")
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HỆ THỐNG QUẢN LÝ KHÁCH SẠN SỬ DỤNG CÔNG NGHỆ WEB\n(BACKEND: SPRING BOOT + FRONTEND: ANGULAR + DATABASE: SQL SERVER)")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Students
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NHÓM SINH VIÊN")
    run.font.size = Pt(13)
    run.bold = True

    students = [
        "NGÔ VÕ TUẤN AN – 24TX810001",
        "TRẦN TRỌNG TÂN – 24TX810025",
        "NGUYỄN KHÔI NGUYÊN – 24TX810011"
    ]
    for s in students:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(s)
        run.font.size = Pt(13)

    doc.add_paragraph()

    # Advisor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GIẢNG VIÊN HƯỚNG DẪN")
    run.font.size = Pt(13)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HUỲNH XUÂN PHỤNG")
    run.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHÓA 2024-2026")
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TP. HỒ CHÍ MINH, 2026.")
    run.font.size = Pt(13)

    doc.add_page_break()


def add_front_matter(doc):
    """Add acknowledgment, declaration, summary."""
    # Lời cảm ơn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LỜI CẢM ƠN")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Trong quá trình thực hiện khóa luận, chúng em đã nhận được sự hướng dẫn, "
        "hỗ trợ và góp ý từ giảng viên, nhà trường, bạn bè và gia đình. Chúng em xin "
        "trân trọng cảm ơn giảng viên hướng dẫn đã dành thời gian định hướng đề tài, "
        "góp ý về phương pháp thực hiện và giúp em nhìn nhận rõ hơn các giới hạn của sản phẩm."
    )
    doc.add_paragraph(
        "Chúng em xin cảm ơn các thầy cô trong khoa đã cung cấp nền tảng kiến thức về "
        "phân tích yêu cầu, thiết kế hệ thống, cơ sở dữ liệu, lập trình và kiểm thử. "
        "Những kiến thức này là cơ sở để em xây dựng, đánh giá và hoàn thiện hệ thống "
        "quản lý khách sạn LuxeStay."
    )
    doc.add_paragraph(
        "Chúng em cũng xin cảm ơn gia đình và bạn bè đã động viên, chia sẻ và hỗ trợ "
        "trong suốt quá trình học tập. Mặc dù đã cố gắng đối chiếu báo cáo với mã nguồn "
        "và bằng chứng kiểm thử hiện hành, khóa luận vẫn có thể còn thiếu sót. Chúng em "
        "mong nhận được nhận xét của quý thầy cô để tiếp tục hoàn thiện sản phẩm và năng "
        "lực chuyên môn."
    )
    doc.add_page_break()

    # Lời cam đoan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LỜI CAM ĐOAN")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Chúng em cam đoan khóa luận về hệ thống quản lý khách sạn LuxeStay là kết quả "
        "học tập, nghiên cứu và triển khai của chúng em dưới sự hướng dẫn của giảng viên "
        "hướng dẫn. Các nội dung, sơ đồ, số liệu kiểm thử và kết luận trong báo cáo "
        "được đối chiếu với mã nguồn, cơ sở dữ liệu, tài liệu kỹ thuật và bằng chứng "
        "tại thời điểm xác minh."
    )
    doc.add_paragraph(
        "Những tài liệu, công nghệ và nguồn tham khảo được sử dụng trong khóa luận sẽ "
        "được trích dẫn trong phần tài liệu tham khảo."
    )
    doc.add_paragraph(
        "Chúng em chịu trách nhiệm về tính trung thực của nội dung báo cáo và cam kết "
        "không cố ý đưa thông tin sai lệch."
    )
    doc.add_page_break()

    # Tóm tắt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÓM TẮT")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Khóa luận xây dựng hệ thống quản lý khách sạn LuxeStay nhằm hỗ trợ hai nhóm "
        "nhu cầu chính: khách hàng tìm kiếm, đặt và theo dõi dịch vụ lưu trú; đơn vị "
        "vận hành quản lý cơ sở, loại phòng, phòng vật lý, đặt phòng và vòng đời lưu trú "
        "trong môi trường nhiều cơ sở. Hệ thống sử dụng Angular cho giao diện, Spring Boot "
        "cho backend và SQL Server cho lưu trữ dữ liệu; Flyway quản lý thay đổi schema."
    )
    doc.add_paragraph(
        "Phiên bản hiện tại hỗ trợ tìm kiếm địa điểm và cơ sở bằng tiếng Việt, xem "
        "RoomType và tồn phòng, đặt một RoomType với số lượng nhiều phòng, thanh toán "
        "VNPay hoặc simulator, hủy booking và ghi nhận hoàn tiền idempotent, xem hóa đơn, "
        "gán phòng, check-in, thêm dịch vụ trong thời gian lưu trú, check-out và tạo tác "
        "vụ dọn phòng."
    )
    doc.add_paragraph(
        "Từ khóa: quản lý khách sạn, đặt phòng, multi-property, RBAC, tồn phòng, "
        "Spring Boot, Angular, SQL Server."
    )
    doc.add_page_break()


def add_screenshot(doc, filename, caption, figure_num):
    """Add a screenshot image with caption if file exists."""
    img_path = SCREENSHOT_DIR / f"{filename}.png"
    if img_path.exists():
        try:
            doc.add_picture(str(img_path), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            p = doc.add_paragraph(f"[Ảnh chưa có: {filename}.png]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Ảnh chưa chụp: {filename}.png - Chạy capture_screenshots.py]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Hình {figure_num}. {caption}")
    run.font.size = Pt(11)
    run.italic = True


def add_table_from_data(doc, headers, rows, table_num=None, caption=None):
    """Add a formatted table to the document."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(f"Bảng {table_num}. {caption}")
        run.bold = True
        run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacer


def main():
    print("📄 Building LuxeStay Thesis DOCX...")

    doc = setup_document()

    # 1. Cover page
    print("  Adding cover page...")
    add_cover_page(doc)

    # 2. Front matter
    print("  Adding front matter...")
    add_front_matter(doc)

    # 3. [Placeholder] Mục lục - Word sẽ tự tạo từ Heading styles
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MỤC LỤC")
    run.font.size = Pt(16)
    run.bold = True
    doc.add_paragraph("[Nhấn Ctrl+A > F9 trong Word để cập nhật mục lục tự động]")
    doc.add_page_break()

    # 4. Danh mục viết tắt
    print("  Adding abbreviation list...")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DANH MỤC TỪ VIẾT TẮT")
    run.font.size = Pt(16)
    run.bold = True
    doc.add_paragraph()

    abbreviations = [
        ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("CRUD", "Create, Read, Update, Delete", "Tạo, Đọc, Cập nhật, Xóa"),
        ("DTO", "Data Transfer Object", "Đối tượng truyền dữ liệu"),
        ("ERD", "Entity Relationship Diagram", "Sơ đồ quan hệ thực thể"),
        ("JWT", "JSON Web Token", "Mã thông báo web JSON"),
        ("RBAC", "Role-Based Access Control", "Kiểm soát truy cập theo vai trò"),
        ("REST", "Representational State Transfer", "Chuyển đổi trạng thái đại diện"),
        ("UI", "User Interface", "Giao diện người dùng"),
        ("UML", "Unified Modeling Language", "Ngôn ngữ mô hình hợp nhất"),
    ]
    add_table_from_data(doc,
        ["Viết tắt", "Tiếng Anh", "Tiếng Việt"],
        abbreviations
    )
    doc.add_page_break()

    # The main body chapters are already in the existing DOCX.
    # This script creates the structural framework with proper formatting.
    # Content from THESIS.md should be pasted into these sections.

    print("  Adding chapter headings and structure...")

    # Chapter 1
    doc.add_heading("CHƯƠNG 1\nTỔNG QUAN ĐỀ TÀI", level=1)
    doc.add_paragraph("[Nội dung Chương 1 - Sao chép từ bản thảo hiện tại hoặc docs/THESIS.md]")
    doc.add_page_break()

    # Chapter 2
    doc.add_heading("CHƯƠNG 2\nCƠ SỞ LÝ THUYẾT", level=1)
    doc.add_paragraph("[Nội dung Chương 2 - Sao chép từ bản thảo hiện tại hoặc docs/THESIS.md]")
    doc.add_page_break()

    # Chapter 3
    doc.add_heading("CHƯƠNG 3\nPHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1)
    doc.add_paragraph("[Nội dung Chương 3 - Sao chép từ bản thảo hiện tại hoặc docs/THESIS.md]")
    doc.add_paragraph("[Chèn các diagram SVG/PNG từ docs/thesis-assets/diagrams/]")
    doc.add_page_break()

    # Chapter 4 - With screenshots
    doc.add_heading("CHƯƠNG 4\nCÀI ĐẶT VÀ KIỂM THỬ HỆ THỐNG", level=1)

    doc.add_heading("4.1. CẤU TRÚC CÀI ĐẶT", level=2)
    doc.add_paragraph(
        "Backend được tổ chức theo các nhóm controllers, services, repositories, entities, "
        "dtos và security. Controller không chứa nghiệp vụ phức tạp; Service điều phối giao "
        "dịch và kiểm tra quy tắc; Repository đóng gói truy vấn dữ liệu."
    )
    doc.add_paragraph(
        "Frontend tổ chức theo core, shared và features. Core chứa dịch vụ dùng toàn ứng "
        "dụng, interceptor và guard. Shared chứa thành phần trình bày dùng lại. Features "
        "chứa màn hình theo nghiệp vụ."
    )

    doc.add_heading("4.2. CÀI ĐẶT XÁC THỰC VÀ PHÂN QUYỀN", level=2)
    doc.add_paragraph(
        "Endpoint đăng nhập phát hành JWT sau khi kiểm tra thông tin tài khoản. Bộ lọc bảo "
        "mật đọc token và tạo Authentication. Annotation @Permission kiểm tra chức năng và "
        "hành động. Các endpoint nhạy cảm còn dùng @PreAuthorize để giới hạn role."
    )

    doc.add_heading("4.3. CÀI ĐẶT TÌM KIẾM VÀ ĐẶT PHÒNG", level=2)
    doc.add_paragraph(
        "Trang chủ cung cấp autocomplete theo địa điểm và cơ sở. Trang kết quả nhận bộ lọc, "
        "sắp xếp và phân trang từ URL hoặc Search State."
    )

    doc.add_heading("4.4. CÀI ĐẶT THANH TOÁN, HỦY VÀ HOÀN TIỀN", level=2)
    doc.add_paragraph(
        "Hệ thống hỗ trợ tạo payment, URL VNPay, callback VNPay và callback simulator. "
        "Callback chỉ ghi nhận thành công khi mã giao dịch chưa tồn tại."
    )

    doc.add_heading("4.5. CÀI ĐẶT VẬN HÀNH LƯU TRÚ", level=2)
    doc.add_paragraph(
        "Nhân viên có thể xem phòng còn trống, gán nhiều phòng vật lý và thực hiện check-in. "
        "Check-out tổng hợp chi phí, tạo hóa đơn, cập nhật phòng thành DIRTY và tạo "
        "housekeeping task."
    )

    doc.add_heading("4.6. CÀI ĐẶT MULTI-PROPERTY VÀ FEATURE GATE", level=2)
    doc.add_paragraph(
        "Active Property Context xác định cơ sở đang được quản lý. Feature Gate kiểm tra "
        "AccountSubscription và giới hạn của gói trước thao tác tạo tài nguyên."
    )

    # 4.7 - Screenshots
    doc.add_heading("4.7. GIAO DIỆN ĐÃ CÀI ĐẶT", level=2)

    doc.add_heading("4.7.1. Giao diện Public", level=3)
    doc.add_paragraph(
        "Trang chủ là điểm đầu tiên khách hàng tiếp cận hệ thống. Trang cung cấp thanh "
        "tìm kiếm nhanh theo địa điểm, ngày và số khách."
    )

    screenshots_public = [
        ("4-01-trang-chu-desktop", "Trang chủ LuxeStay trên desktop", "4.1"),
        ("4-02-trang-chu-mobile", "Trang chủ LuxeStay trên mobile", "4.2"),
        ("4-03-autocomplete-tim-kiem", "Autocomplete tìm kiếm địa điểm", "4.3"),
        ("4-04-ket-qua-tim-kiem", "Kết quả tìm kiếm với bộ lọc và phân trang", "4.4"),
        ("4-05-chi-tiet-khach-san", "Chi tiết khách sạn", "4.5"),
        ("4-06-chon-loai-phong", "Chọn loại phòng và số lượng phòng", "4.6"),
    ]
    for fname, caption, fig_num in screenshots_public:
        add_screenshot(doc, fname, caption, fig_num)

    doc.add_heading("4.7.2. Giao diện Xác thực", level=3)
    add_screenshot(doc, "4-07-dang-ky-tai-khoan", "Đăng ký tài khoản", "4.7")
    add_screenshot(doc, "4-08-dang-nhap", "Đăng nhập", "4.8")

    doc.add_heading("4.7.3. Giao diện Đặt phòng và Thanh toán", level=3)
    screenshots_booking = [
        ("4-09-checkout-dat-phong", "Trang checkout đặt phòng", "4.9"),
        ("4-12-danh-sach-booking", "Danh sách booking của khách hàng", "4.10"),
        ("4-13-lich-su-hoan-tien", "Lịch sử hoàn tiền", "4.11"),
        ("4-14-hoa-don-khach-hang", "Hóa đơn khách hàng", "4.12"),
        ("4-15-trang-ca-nhan", "Trang cá nhân", "4.13"),
    ]
    for fname, caption, fig_num in screenshots_booking:
        add_screenshot(doc, fname, caption, fig_num)

    doc.add_heading("4.7.4. Giao diện Quản trị", level=3)
    doc.add_paragraph(
        "Khu vực quản trị dùng sidebar điều hướng, menu được tạo từ dữ liệu quyền API."
    )
    screenshots_admin = [
        ("4-17-dashboard-tong-quan", "Dashboard tổng quan", "4.14"),
        ("4-18-quan-ly-nguoi-dung", "Quản lý người dùng", "4.15"),
        ("4-20-quan-ly-vai-tro", "Quản lý vai trò", "4.16"),
        ("4-21-phan-quyen-vai-tro", "Ma trận phân quyền", "4.17"),
        ("4-22-quan-ly-co-so", "Quản lý cơ sở lưu trú", "4.18"),
        ("4-23-quan-ly-loai-phong", "Quản lý loại phòng", "4.19"),
        ("4-24-quan-ly-phong-vat-ly", "Quản lý phòng vật lý", "4.20"),
        ("4-25-quan-ly-dich-vu", "Quản lý dịch vụ khách sạn", "4.21"),
    ]
    for fname, caption, fig_num in screenshots_admin:
        add_screenshot(doc, fname, caption, fig_num)

    doc.add_heading("4.7.5. Giao diện Vận hành lưu trú", level=3)
    screenshots_ops = [
        ("4-26-quan-ly-dat-phong", "Quản lý đặt phòng", "4.22"),
        ("4-27-timeline-dat-phong", "Timeline đặt phòng", "4.23"),
        ("4-29-quan-ly-hoa-don", "Quản lý hóa đơn", "4.24"),
        ("4-43-housekeeping", "Quản lý Housekeeping", "4.25"),
    ]
    for fname, caption, fig_num in screenshots_ops:
        add_screenshot(doc, fname, caption, fig_num)

    doc.add_heading("4.7.6. Giao diện Hệ thống", level=3)
    screenshots_system = [
        ("4-30-import-co-so", "Import cơ sở từ dữ liệu mở", "4.26"),
        ("4-31-property-claims", "Property Claims", "4.27"),
        ("4-32-subscription-plans", "Subscription Plans", "4.28"),
        ("4-33-chat-ho-tro", "Chat hỗ trợ trung tâm", "4.29"),
        ("4-34-audit-log", "Audit Log", "4.30"),
    ]
    for fname, caption, fig_num in screenshots_system:
        add_screenshot(doc, fname, caption, fig_num)

    doc.add_heading("4.7.7. Giao diện Management Portal", level=3)
    screenshots_mgmt = [
        ("4-40-management-dashboard", "Management Dashboard", "4.31"),
        ("4-46-property-revenue", "Báo cáo doanh thu cơ sở", "4.32"),
        ("4-47-subscription-billing", "Subscription Billing", "4.33"),
    ]
    for fname, caption, fig_num in screenshots_mgmt:
        add_screenshot(doc, fname, caption, fig_num)

    # 4.8 - Test Results
    doc.add_heading("4.8. CHIẾN LƯỢC KIỂM THỬ", level=2)
    doc.add_paragraph(
        "Kiểm thử được chia thành: Unit test kiểm tra Service với dependency giả lập; "
        "Integration test khởi tạo Spring context, MockMvc và H2; Frontend unit test "
        "kiểm tra component và service; E2E test Playwright chạy các luồng public, "
        "customer, payment và admin."
    )

    doc.add_heading("4.9. KẾT QUẢ KIỂM THỬ", level=2)

    # Test table - Authentication
    add_table_from_data(doc,
        ["STT", "Kịch bản", "Đầu vào", "Kết quả mong đợi", "Đánh giá"],
        [
            ["1", "Đăng ký thành công", "Email, mật khẩu hợp lệ", "Tạo tài khoản, trả JWT", "✅"],
            ["2", "Đăng ký email trùng", "Email đã tồn tại", "HTTP 400", "✅"],
            ["3", "Đăng nhập thành công", "Đúng email/password", "JWT + user info", "✅"],
            ["4", "Đăng nhập sai mật khẩu", "Sai password", "HTTP 401", "✅"],
            ["5", "Token hết hạn", "JWT expired", "HTTP 401", "✅"],
            ["6", "Truy cập không có quyền", "User thiếu ROLE", "HTTP 403", "✅"],
            ["7", "Action Mask kiểm tra", "Role thiếu DELETE", "API trả 403", "✅"],
        ],
        "4.2", "Kịch bản kiểm thử xác thực và phân quyền"
    )

    # Summary test table
    add_table_from_data(doc,
        ["Hạng mục", "Số test", "Đạt", "Tỷ lệ"],
        [
            ["Backend Unit Test (JUnit)", "123", "123", "100%"],
            ["Frontend Unit Test (Jasmine)", "73", "73", "100%"],
            ["Angular Production Build", "1", "1", "100%"],
            ["Playwright E2E", "5", "2", "40%"],
            ["Kiểm thử thủ công", "36", "36", "100%"],
            ["Tổng cộng", "238", "235", "98.7%"],
        ],
        "4.7", "Tổng hợp kết quả kiểm thử"
    )

    doc.add_heading("4.10. ĐÁNH GIÁ KẾT QUẢ", level=2)
    doc.add_paragraph(
        "Kết quả cài đặt đáp ứng các nghiệp vụ cốt lõi gồm xác thực, tìm kiếm, đặt phòng, "
        "thanh toán, quản lý tồn phòng và vận hành lưu trú. Các kiểm thử backend và frontend "
        "cho thấy những quy tắc nghiệp vụ chính hoạt động ổn định."
    )

    doc.add_page_break()

    # Chapter 5
    doc.add_heading("CHƯƠNG 5\nKẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_paragraph("[Nội dung Chương 5 - Sao chép từ bản thảo hiện tại]")

    doc.add_page_break()

    # References
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        '[1] Angular, "Angular Documentation," https://angular.dev/.',
        '[2] Chart.js, "Chart.js Documentation," https://www.chartjs.org/docs/.',
        '[3] Docker, "Docker Documentation," https://docs.docker.com/.',
        '[4] Flyway, "Flyway Documentation," https://documentation.red-gate.com/flyway/.',
        '[5] Oracle, "Java Platform, Standard Edition Documentation," https://docs.oracle.com/en/java/javase/21/.',
        '[6] Playwright, "Playwright Documentation," https://playwright.dev/docs/intro.',
        '[7] PrimeTek, "PrimeNG Documentation," https://primeng.org/.',
        '[8] Spring, "Spring Boot Reference Documentation," https://docs.spring.io/spring-boot/docs/3.2.5/reference/html/.',
        '[9] Spring, "Spring Security Reference," https://docs.spring.io/spring-security/reference/.',
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"\n✅ Thesis DOCX saved to: {OUTPUT_FILE}")
    print(f"   File size: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")
    print()
    print("📝 Next steps:")
    print("   1. Run capture_screenshots.py to capture screenshots")
    print("   2. Re-run this script to embed screenshots")
    print("   3. Open in Word and insert Table of Contents (Ctrl+A > F9)")
    print("   4. Copy detailed chapter content from THESIS.md")
    print("   5. Insert diagram images from docs/thesis-assets/diagrams/")


if __name__ == "__main__":
    main()
